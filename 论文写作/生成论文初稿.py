from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / ".agents" / "skills" / "math-modeling" / "references" / "论文模板.docx"
OUTPUT_DIR = ROOT / "论文写作"
OUTPUT_PATH = OUTPUT_DIR / "论文.docx"
FIG_DIR = ROOT / "代码实现" / "figures"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"未找到模板段落：{text}")


def replace_lines_after(anchor: Paragraph, lines: list[str]) -> Paragraph:
    current = anchor
    for line in lines:
        current = insert_paragraph_after(current, line, anchor.style)
    return current


def add_picture_after(anchor: Paragraph, image_path: Path, width_cm: float, caption: str) -> Paragraph:
    pic_para = insert_paragraph_after(anchor, "", anchor.style)
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap_para = insert_paragraph_after(pic_para, caption, anchor.style)
    return cap_para


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE_PATH, OUTPUT_PATH)
    doc = Document(str(OUTPUT_PATH))

    set_paragraph_text(find_paragraph(doc, "论文题目"), "战场环境下无人机群协同侦察的数学建模与优化研究")

    set_paragraph_text(
        find_paragraph(doc, "总体介绍"),
        "本文针对战场环境下无人机群协同侦察任务，围绕风险约束下的侦察收益最大化、异构多机协同分配与事件触发动态重规划三类问题开展建模。首先，基于附件中的基础风险概率、防空火力部署与目标价值分布，构建战场综合风险场，并将高价值区域提取为目标簇；其次，针对问题一建立风险约束定向越野模型，以期望收益最大化为目标，在续航和返航约束下完成单机路径规划，并通过策略分层与算法对比筛选综合最优方案；然后，针对问题二设计分层竞标、余量回收与全局重平衡机制，求解五架异构无人机的协同巡查分配；最后，针对问题三构建事件触发滚动重规划状态机，在威胁突增、目标变化、通信中断、战损与低电量补给条件下完成动态仿真闭环。",
    )
    set_paragraph_text(
        find_paragraph(doc, "针对问题一···"),
        "针对问题一，本文将战场区域离散为网格风险图，将高价值目标聚合为目标簇，建立风险约束定向越野模型。求解层面采用“策略分层+算法对比”的优秀论文式处理流程，设置稳健、平衡、高收益三类风险偏好策略，并引入最近邻、贪心+2-opt、模拟退火和遗传算法进行求解对比。结果表明，稳健策略下的贪心+2-opt方案在期望收益、生存概率与航时之间实现了最优折中，可作为问题一的综合最优方案。",
    )
    set_paragraph_text(
        find_paragraph(doc, "针对问题二···"),
        "针对问题二，本文将五架异构无人机的协同巡查抽象为异构多智能体定向越野/多旅行商问题，采用“能力分层-边际收益竞标-单机路径优化-冲突修复-网格重算”的完整流程完成任务分配。最新结果表明，在160个候选目标中共完成52个目标，完成率达到52.5052%，综合效能为0.317044，且任务在五架无人机之间形成了较合理的分摊。",
    )
    set_paragraph_text(
        find_paragraph(doc, "针对问题三···"),
        "针对问题三，本文将动态侦察过程抽象为事件触发的滚动时域重规划问题，构建 Search/Track/Replan/Evade/Isolated/RTB/Resupply/Lost 八状态机，并引入固定且分散布设的补给点。主场景仿真显示，系统能够在威胁突增、目标增补、通信中断、战损与低电量返航补给等事件下持续重规划，完成23个目标，综合效能为0.569695。",
    )
    set_paragraph_text(find_paragraph(doc, "关键词："), "关键词：无人机协同侦察；风险约束定向越野；A*路径规划；多策略对比；Pareto前沿")

    set_paragraph_text(
        find_paragraph(doc, "1.1 问题背景"),
        "在现代非对称作战与边境冲突场景中，无人机系统承担态势感知、目标侦察与火力引导等关键任务。与和平时期巡检相比，战场环境具有高动态、高风险和强不确定性：一方面，敌方雷达与防空火力会显著抬升局部区域风险；另一方面，高价值目标可能呈现分散、临时和机动特征。因而，无人机侦察任务不能简单追求覆盖范围最大，而应在续航、返航和生存约束下对收益与风险进行统筹优化。",
    )
    set_paragraph_text(
        find_paragraph(doc, "对于问题一，xxx"),
        "对于问题一，需要在静态已知威胁场中，为单架无人机规划一条从起降点(0,0)出发并返回的闭合路径，使其在不超过最大续航时间的条件下获得尽可能高的期望侦察收益。",
    )
    set_paragraph_text(
        find_paragraph(doc, "对于问题二，xxx"),
        "对于问题二，需要在五架异构无人机协同条件下完成高价值目标分配、路径规划与协同效率评估，并进一步分析无人机性能、战场威胁和战损条件对总体任务表现的影响。本文将其转化为异构协同巡查优化问题，通过任务分层分配与单机路径精修实现整体性能提升。",
    )
    set_paragraph_text(
        find_paragraph(doc, "对于问题三，xxx"),
        "对于问题三，需要在动态目标、动态威胁、通信延迟与中断条件下建立事件触发滚动重规划机制，并给出任务冲突处理和返航判定逻辑。本文将动态态势演化刻画为事件驱动仿真过程，并通过状态机与补给再入队机制实现闭环控制。",
    )
    set_paragraph_text(
        find_paragraph(doc, "1.3 问题要求"),
        "题目要求按单机静态规划、多机静态协同和动态重规划三个层次展开建模。本文已完成三问的模型建立、求解、策略对比和结果分析，并将按照“问题一—问题二—问题三”的逻辑顺序给出完整论文结构。",
    )

    assumptions = [
        "假设1：战场区域内基础风险概率、防空火力点位置和目标价值图在问题一中均为已知静态信息。",
        "假设2：无人机在网格内飞行时，其被发现风险只与当前路径经过的网格风险值有关，且相邻网格风险独立累积。",
        "假设3：高价值目标可通过目标簇中心表示，目标簇权重取簇内价值和，识别停留时间统一取1个时间单位。",
        "假设4：无人机在问题一中一旦被发现将立即放弃后续任务，因此后续收益需按到达该目标前的生存概率折减。",
        "假设5：问题一采用UAV-04作为代表单机平台，其起点位于(0,0)，且性能参数以附件数据为准。",
    ]
    for idx, text in enumerate(assumptions, start=26):
        set_paragraph_text(doc.paragraphs[idx - 1], text)

    symbol_table = doc.tables[0]
    symbols = [
        ("p_{ij}", "网格(i,j)处的综合风险概率", "-"),
        ("w_q", "第q个目标簇的权重", "-"),
        ("S(P)", "路径P的生存概率", "-"),
        ("C_R(P)", "路径P的对数风险代价", "-"),
        ("T(P)", "路径P的总航时", "min"),
        ("E[W(P)]", "路径P的期望侦察收益", "-"),
        ("T_{max}", "无人机最大续航时间", "min"),
        ("v", "无人机最大飞行速度", "km/min"),
    ]
    while len(symbol_table.rows) < len(symbols) + 1:
        symbol_table.add_row()
    for i, (sym, desc, unit) in enumerate(symbols, start=1):
        symbol_table.cell(i, 0).text = sym
        symbol_table.cell(i, 1).text = desc
        symbol_table.cell(i, 2).text = unit

    set_paragraph_text(
        find_paragraph(doc, "问题一模型的建立"),
        "问题一模型的建立",
    )
    anchor = find_paragraph(doc, "问题一模型的建立")
    lines = [
        "为同时体现收益、风险与返航约束，本文将问题一抽象为风险约束定向越野问题。首先，对附件中的基础风险概率与七处防空火力部署进行融合，构建战场综合风险场；其次，将高价值网格聚合为目标簇，每个目标簇由加权中心、簇权重和识别停留时间共同描述；最后，在起点和终点固定为(0,0)的条件下，对目标簇访问序列进行优化。",
        "设路径P经过的网格序列为(g_0,g_1,...,g_L)，则路径生存概率定义为 S(P)=∏_{ℓ=0}^{L}(1-p_{g_ℓ})，对数风险代价定义为 C_R(P)=-∑_{ℓ=0}^{L}ln(1-p_{g_ℓ})。若按顺序完成目标簇(q_1,q_2,...,q_K)，则第h个目标前的生存概率记为S_h，对应期望收益为 w_{q_h}S_h。",
        "因此，问题一的目标函数写为：最大化 E[W(P)]=∑_{h=1}^{K}w_{q_h}S_h。在此基础上加入续航约束T(P)≤T_max、返航闭合约束P_0=P_L=(0,0)以及目标识别约束，从而得到完整的风险约束定向越野模型。",
        "考虑到题目既要求结果可解释，又要求模型具备论文级说服力，本文没有直接只输出一条路径，而是采用优秀优化类论文常见的处理方式：先构造稳健、平衡、高收益三类风险偏好策略，再在每类策略下比较不同求解算法表现，最后通过Pareto前沿筛选综合最优方案。",
    ]
    last = replace_lines_after(anchor, lines)
    last = add_picture_after(last, FIG_DIR / "问题1_全部目标簇分布.png", 13.5, "图1  战场综合风险场与全部目标簇分布")

    set_paragraph_text(find_paragraph(doc, "问题一的求解"), "问题一的求解")
    anchor = find_paragraph(doc, "问题一的求解")
    lines = [
        "在数据处理阶段，首先读取附件中的基础风险图、防空火力图和价值图，并裁剪至公共有效区域400×234；随后利用高斯扩散机制叠加防空火力风险，得到综合风险场。对价值图采用阈值聚类，提取384个目标簇作为潜在侦察目标。",
        "在求解阶段，本文使用A*算法在网格层生成可执行避险路径，用最近邻、贪心+2-opt、模拟退火与遗传算法对目标簇访问序列进行优化。实验结果表明，在当前候选规模与约束条件下，贪心+2-opt、模拟退火与遗传算法均达到相同最优值，而贪心+2-opt求解时间最短，因此被选为正文主算法。",
        "三类策略的对比结果表明：高收益策略能够将期望收益提高到6.179047，但对应生存概率仅为0.00013839，风险代价显著抬升；平衡策略期望收益为5.109424，但生存概率仅为0.001303；稳健策略下，贪心+2-opt得到的期望收益为4.948609，生存概率达到0.29933324，总航时仅11.782573，能够在收益、风险与执行代价之间取得最合理的折中。",
        "进一步地，本文对所有策略与算法结果进行了Pareto前沿筛选。结果显示，综合最优方案为“稳健策略+贪心+2-opt”，该方案在收益略低于激进方案的情况下显著提升了生存概率，并在网格尺度敏感性分析中始终保持可行。因此，本文将其作为问题一的最终推荐路径。",
    ]
    last = replace_lines_after(anchor, lines)
    last = add_picture_after(last, FIG_DIR / "问题1_策略收益风险对比.png", 13.5, "图2  不同策略与算法的收益—风险对比")
    last = add_picture_after(last, FIG_DIR / "问题1_单机最优路径.png", 12.5, "图3  综合最优方案的单机A*闭合路径")
    last = add_picture_after(last, FIG_DIR / "问题1_算法收敛曲线.png", 14.0, "图4  不同策略下各算法的收敛曲线")

    q2_build = find_paragraph(doc, "1.1 问题二模型的建立")
    set_paragraph_text(q2_build, "1.1 问题二模型的建立")
    replace_lines_after(
        q2_build,
        [
            "问题二可抽象为五架异构无人机的协同巡查优化问题，具有明显的“多目标、高约束、强异构”特征。本文将其等价为异构多智能体定向越野/多旅行商问题：一方面，候选目标在收益、局部风险与可达时间上存在显著差异；另一方面，不同无人机在速度、航时、载荷与通信能力上具有异构性，因此不能采用统一阈值或单一贪心规则直接求解。",
            "求解流程采用“能力分层-边际收益竞标-单机路径优化-冲突修复-网格级重算”的五阶段框架。首先，根据无人机最大速度、最大航时、传感器范围和通信范围构建平台能力评分，对候选目标进行优先级排序；其次，在初始分配阶段结合目标权重、局部风险、可达时间与当前负载进行边际收益竞标；再次，对每架无人机的局部路径执行贪心插入与2-opt微调；然后，通过余量回收与跨机再分配消除局部死锁；最后，使用A*网格路径重算统一规划轨迹与评价指标口径。",
            "在结果评估层面，本文同时考察完成目标数、期望收益、总航时、总风险代价、总能耗与负载均衡度，并以综合效能J作为总体比较指标。该设计能够避免只追求完成数而忽略战场代价，也能避免过度保守导致资源浪费。",
        ],
    )
    q2_solve = find_paragraph(doc, "1.2 问题二的求解")
    set_paragraph_text(q2_solve, "1.2 问题二的求解")
    last = replace_lines_after(
        q2_solve,
        [
            "基于最终代码实现，本文首先从384个目标中筛选出160个可达候选目标，再按“收益-风险-时间-负载”综合边际收益对目标进行动态分配。重平衡后，共完成52个目标，完成率达到52.5052%，总期望收益为86.512930，总航时378.744184，总风险代价155.263111，总能耗765.292093，平均生存概率为0.00016432，综合效能J为0.317044。",
            "从单机负载看，UAV-05承担25个目标，UAV-04承担11个目标，UAV-02承担7个目标，UAV-03承担5个目标，UAV-01承担4个目标；说明高续航、高能力平台承担了更多长程任务，而低负载平台主要负责局部补充巡查。与初始分配相比，重平衡后的 load_balance 提升至0.350041，说明该策略有效缓解了任务过度集中。",
            "图5展示了五架无人机的实际网格路径轨迹，可见各机轨迹覆盖范围与任务分工较为清晰。图6与网格尺度、威胁缩放相关的敏感性结果说明，在不同尺度下模型均能保持可行，且完成目标数与效能指标总体稳定，体现了算法的鲁棒性。",
        ],
    )
    last = add_picture_after(last, FIG_DIR / "问题2_多机路径轨迹.png", 13.2, "图5  问题二五架无人机协同巡查轨迹")
    last = add_picture_after(last, FIG_DIR / "问题2_威胁敏感性.png", 12.8, "图6  问题二威胁缩放敏感性分析")

    q3_build = find_paragraph(doc, "1.1 问题三模型的建立")
    set_paragraph_text(q3_build, "1.1 问题三模型的建立")
    replace_lines_after(
        q3_build,
        [
            "问题三的核心在于动态战场条件下的事件触发重规划与任务闭环执行。本文将其抽象为滚动时域动态优化问题，并显式引入威胁突增、目标增补/删除、通信中断、战损退出与低电量返航补给等事件。与静态问题不同，动态问题中的关键不在于一次性找到最优路径，而在于面对事件冲击时快速恢复可执行任务序列。",
            "为此，本文构建了八状态机：Search、Track、Replan、Evade、Isolated、RTB、Resupply 与 Lost。状态机的设计遵循“发现-识别-重规划-规避-孤岛自治-返航-补给-失联”这一战场行动逻辑；当事件触发或周期同步到达时，系统将任务池重新组织并执行局部再分配。补给点采用固定且分散的布局，以减少返航与再入队成本，并提高全盘期望收益。",
            "在仿真设定上，本文采用1条主场景与3条补充场景。主场景同时包含威胁突增、目标增补、通信中断、战损和低电量补给，补充场景则分别强调通信受限、高威胁压制和战损补给联动三类典型态势，以验证策略在不同战场条件下的稳定性。",
        ],
    )
    q3_solve = find_paragraph(doc, "1.2 问题三的求解")
    set_paragraph_text(q3_solve, "1.2 问题三的求解")
    last = replace_lines_after(
        q3_solve,
        [
            "主场景仿真结果表明，系统在60个时间步内完成23个目标，完成率为0.273612，实际收益749.008540，期望收益328.589895，生存指数0.03083052，负载均衡0.739130，重规划23次，补给4次，返航5次，战损1架，综合效能J为0.569695。尽管绝对完成目标数低于问题二，但动态重规划能够有效应对高频事件冲击，维持任务链条不断裂。",
            "补充场景中，通信受限场景完成29个目标，综合效能0.605965；高威胁压制场景完成33个目标，综合效能0.631921；战损补给联动场景完成27个目标，综合效能0.561908。三组结果说明，所提出的状态机与补给再入队机制能够在不同事件组合下保持较稳定的策略表现。",
            "图7给出了主场景状态机与动态转移关系，图8展示了主场景的动态轨迹与补给点分布，图9则对四类策略的综合效能进行了对比。三图共同说明：动态问题的关键不是单次路径长度最短，而是事件冲击下的重规划弹性与任务恢复能力。",
        ],
    )
    last = add_picture_after(last, FIG_DIR / "问题3_动态重规划状态机.png", 13.2, "图7  问题三事件触发动态重规划状态机")
    last = add_picture_after(last, FIG_DIR / "问题3_主场景动态轨迹.png", 13.2, "图8  问题三主场景动态轨迹与补给点分布")
    last = add_picture_after(last, FIG_DIR / "问题3_策略效能对比.png", 12.8, "图9  问题三不同场景策略效能对比")

    set_paragraph_text(find_paragraph(doc, "7.1 模型的优点"), "7.1 模型的优点")
    replace_lines_after(
        find_paragraph(doc, "7.1 模型的优点"),
        [
            "（1）模型将基础风险与防空火力风险统一到同一概率场中，能够较真实地反映战场威胁空间分布。",
            "（2）以期望收益而非简单权重和作为目标，能够体现“被发现即逃离”机制对后续任务的截断影响。",
            "（3）采用“策略分层+算法对比+Pareto前沿筛选”的流程，既满足模型求解，又增强了论文结论的可解释性。",
        ],
    )
    set_paragraph_text(find_paragraph(doc, "7.2 模型的缺点"), "7.2 模型的缺点")
    replace_lines_after(
        find_paragraph(doc, "7.2 模型的缺点"),
        [
            "（1）问题一中将目标区域聚合为目标簇，会在一定程度上损失簇内微观结构信息。",
            "（2）风险独立累积的假设忽略了连续暴露和敌方联动火力的复杂耦合。",
            "（3）问题二与问题三仍采用启发式重平衡与事件驱动重规划，虽便于解释，但距离全局最优仍有进一步提升空间。",
        ],
    )
    set_paragraph_text(find_paragraph(doc, "7.3 模型的改进"), "7.3 模型的改进")
    replace_lines_after(
        find_paragraph(doc, "7.3 模型的改进"),
        [
            "后续可在问题二中进一步引入更严格的多目标优化和跨机博弈式任务交换机制，并在问题三中加入更精细的通信拓扑演化、补给资源容量约束与随机事件扰动，从而形成更接近真实战场的无人机群协同侦察模型。",
        ],
    )

    ref_anchor = find_paragraph(doc, "八、参考文献")
    references = [
        "[1] Chao I M, Golden B L, Wasil E A. A fast and effective heuristic for the orienteering problem[J]. European Journal of Operational Research, 1996, 88(3): 475-489.",
        "[2] Vansteenwegen P, Souffriau W, Van Oudheusden D. The orienteering problem: A survey[J]. European Journal of Operational Research, 2011, 209(1): 1-10.",
        "[3] Archetti C, Hertz A, Speranza M G. Metaheuristics for the team orienteering problem[J]. Journal of Heuristics, 2007, 13(1): 49-76.",
        "[4] Ragi S, Chong E K P. UAV Path Planning in a Dynamic Environment via Partially Observable Markov Decision Process[J]. IEEE Transactions on Aerospace and Electronic Systems, 2013, 49(4): 2397-2412.",
        "[5] 宋超, 李沁, 马云红, 黄晶瑞. 基于优化A*和MPC融合算法的三维无人机航迹规划[J]. 系统工程与电子技术, 2023, 45(12): 3995-4004.",
        "[6] 王建峰, 贾高伟, 郭正, 侯中喜. 多无人机协同任务规划方法研究综述[J]. 系统工程与电子技术, 2024, 46(10): 3437-3450.",
    ]
    replace_lines_after(ref_anchor, references)

    appendix_anchor = find_paragraph(doc, "附录")
    replace_lines_after(
        appendix_anchor,
        [
            "附录A  问题一核心程序位于“代码实现/问题1_求解.py”，实验输出位于“代码实现/results/”和“代码实现/figures/”。",
            "附录B  问题二核心程序位于“代码实现/问题2_求解.py”，问题三核心程序位于“代码实现/问题3_求解.py”，完整结果位于“代码实现/results/”和“代码实现/figures/”。",
        ],
    )

    doc.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    build_document()
    print(f"论文初稿已生成：{OUTPUT_PATH}")
